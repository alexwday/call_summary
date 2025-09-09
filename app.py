"""
Flask web application for document chatbot with integrated STT and TTS.
"""

import os
import uuid
from datetime import datetime
from flask import Flask, render_template, request, jsonify, Response, session
from flask_cors import CORS
from werkzeug.utils import secure_filename
import docx
import PyPDF2
from pathlib import Path
from src.call_summary.main import model
from src.call_summary.utils.logging import get_logger
from src.call_summary.utils.settings import config
import tempfile
import time
import io
import soundfile as sf
import numpy as np
import ssl
import urllib.request

# Configure SSL for model downloads
def configure_ssl():
    """Configure SSL settings based on environment."""
    ssl_verify = os.environ.get('SSL_VERIFY', 'true').lower() == 'true'
    ssl_cert_file = os.environ.get('SSL_CERT_FILE', 'rbc-ca-bundle.cer')
    
    if ssl_verify and os.path.exists(ssl_cert_file):
        print(f"🔒 Using SSL certificate: {ssl_cert_file}")
        # Create SSL context with corporate certificate
        ssl_context = ssl.create_default_context(cafile=ssl_cert_file)
        # Set as default for urllib
        urllib.request.install_opener(urllib.request.build_opener(
            urllib.request.HTTPSHandler(context=ssl_context)))
        # Set for requests library
        os.environ['REQUESTS_CA_BUNDLE'] = ssl_cert_file
        return ssl_context
    elif not ssl_verify:
        print("⚠️  SSL verification disabled (development mode)")
        # Create unverified context
        ssl_context = ssl.create_default_context()
        ssl_context.check_hostname = False
        ssl_context.verify_mode = ssl.CERT_NONE
        # Set as default for urllib
        urllib.request.install_opener(urllib.request.build_opener(
            urllib.request.HTTPSHandler(context=ssl_context)))
        # Disable for requests
        os.environ['CURL_CA_BUNDLE'] = ""
        os.environ['REQUESTS_CA_BUNDLE'] = ""
        return ssl_context
    else:
        print("✓ Using system default SSL settings")
        return None

# Configure SSL before importing ML models
ssl_context = configure_ssl()

# Import ML models for voice capabilities
import mlx_whisper
from mlx_audio.tts.models.kokoro import KokoroPipeline
from mlx_audio.tts.utils import load_model

logger = get_logger()

app = Flask(__name__)
CORS(app)
app.secret_key = os.environ.get('SECRET_KEY', 'dev-secret-key-change-in-production')
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max file size
app.config['UPLOAD_FOLDER'] = 'uploads'

# Map of Whisper model sizes to their MLX paths
WHISPER_MODELS = {
    'tiny': 'mlx-community/whisper-tiny-mlx',
    'base': 'mlx-community/whisper-base-mlx',
    'small': 'mlx-community/whisper-small-mlx',
    'medium': 'mlx-community/whisper-medium-mlx',
    'large': 'mlx-community/whisper-large-v3-mlx'
}

# Default Whisper model - using small for good balance of speed and accuracy
default_whisper_model = 'small'
whisper_model_path = WHISPER_MODELS[default_whisper_model]

# Initialize voice models
print("Loading voice models...")
print(f"Loading Whisper models (default: {default_whisper_model})...")
print("Whisper model ready!")

print("Loading Kokoro 82M model...")
tts_model_id = 'mlx-community/Kokoro-82M-4bit'
tts_model = load_model(tts_model_id)
tts_pipeline = KokoroPipeline(lang_code='a', model=tts_model, repo_id=tts_model_id)
print("Kokoro TTS model ready!")
print("All voice models loaded successfully!")

# Allowed file extensions
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc', 'txt'}

# Create upload folder if it doesn't exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Store documents in session (in production, use a database)
SESSIONS = {}

# Model options from config
MODEL_OPTIONS = {
    'small': {
        'name': config.llm.small.model,
        'cost_input': config.llm.small.cost_per_1k_input,
        'cost_output': config.llm.small.cost_per_1k_output
    },
    'medium': {
        'name': config.llm.medium.model,
        'cost_input': config.llm.medium.cost_per_1k_input,
        'cost_output': config.llm.medium.cost_per_1k_output
    },
    'large': {
        'name': config.llm.large.model,
        'cost_input': config.llm.large.cost_per_1k_input,
        'cost_output': config.llm.large.cost_per_1k_output
    }
}


def allowed_file(filename):
    """Check if file extension is allowed."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text_from_pdf(file_path):
    """Extract text from PDF file."""
    text = []
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text.append(page.extract_text())
        return '\n'.join(text)
    except Exception as e:
        logger.error(f"Error extracting PDF text: {e}")
        return f"Error reading PDF: {str(e)}"


def extract_text_from_docx(file_path):
    """Extract text from Word document."""
    try:
        doc = docx.Document(file_path)
        paragraphs = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    paragraphs.append(' | '.join(row_text))
        
        return '\n'.join(paragraphs)
    except Exception as e:
        logger.error(f"Error extracting DOCX text: {e}")
        return f"Error reading Word document: {str(e)}"


def extract_text_from_file(file_path, filename):
    """Extract text based on file type."""
    extension = filename.rsplit('.', 1)[1].lower()
    
    if extension == 'pdf':
        return extract_text_from_pdf(file_path)
    elif extension in ['docx', 'doc']:
        return extract_text_from_docx(file_path)
    elif extension == 'txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    else:
        return "Unsupported file type"


def get_file_metadata(file_path, original_filename):
    """Extract comprehensive file metadata."""
    path = Path(file_path)
    stat = path.stat()
    
    metadata = {
        'original_filename': original_filename,
        'file_extension': original_filename.rsplit('.', 1)[1].lower() if '.' in original_filename else 'unknown',
        'file_size_bytes': stat.st_size,
        'file_size_human': format_file_size(stat.st_size),
        'upload_timestamp': datetime.now().isoformat(),
        'last_modified': datetime.fromtimestamp(stat.st_mtime).isoformat(),
        'file_path': str(path.absolute())
    }
    
    return metadata


def format_file_size(size_bytes):
    """Format file size in human-readable format."""
    for unit in ['B', 'KB', 'MB', 'GB']:
        if size_bytes < 1024.0:
            return f"{size_bytes:.2f} {unit}"
        size_bytes /= 1024.0
    return f"{size_bytes:.2f} TB"


@app.route('/')
def index():
    """Main chat interface."""
    if 'session_id' not in session:
        session['session_id'] = str(uuid.uuid4())
        SESSIONS[session['session_id']] = {
            'documents': [],
            'messages': [],
            'total_tokens': {'input': 0, 'output': 0},
            'total_cost': 0.0,
            'selected_model': 'large',  # Default model
            'prompt_mode': 'stage1'  # Default prompt mode
        }
    
    return render_template('chat.html')


@app.route('/upload', methods=['POST'])
def upload_document():
    """Handle document upload."""
    try:
        if 'session_id' not in session:
            return jsonify({'error': 'No session found'}), 400
        
        session_id = session['session_id']
        
        if session_id not in SESSIONS:
            SESSIONS[session_id] = {
                'documents': [],
                'messages': [],
                'total_tokens': {'input': 0, 'output': 0},
                'total_cost': 0.0,
                'selected_model': 'large',  # Default model
                'prompt_mode': 'stage1'  # Default prompt mode
            }
        
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if not allowed_file(file.filename):
            return jsonify({'error': 'File type not allowed. Please upload PDF or Word documents.'}), 400
        
        # Save file
        filename = secure_filename(file.filename)
        unique_filename = f"{session_id}_{uuid.uuid4().hex[:8]}_{filename}"
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
        file.save(file_path)
        
        # Extract text
        text_content = extract_text_from_file(file_path, filename)
        
        # Get comprehensive file metadata
        metadata = get_file_metadata(file_path, filename)
        
        # Store document content with unique ID and metadata
        doc_id = str(uuid.uuid4())[:8]
        SESSIONS[session_id]['documents'].append({
            'id': doc_id,
            'filename': filename,
            'content': text_content,
            'path': file_path,
            'metadata': metadata
        })
        
        # Clean up file after extraction (optional)
        # os.remove(file_path)
        
        logger.info(f"Document uploaded: {filename} for session {session_id}")
        
        # Calculate token count (approximate)
        token_count = len(text_content.split())
        file_size = os.path.getsize(file_path)
        
        return jsonify({
            'success': True,
            'id': doc_id,
            'filename': filename,
            'size': file_size,
            'token_count': token_count,
            'message': f'Successfully uploaded {filename}',
            'document_count': len(SESSIONS[session_id]['documents'])
        })
        
    except Exception as e:
        logger.error(f"Upload error: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/remove_document/<doc_id>', methods=['DELETE'])
def remove_document(doc_id):
    """Remove a document from the session."""
    try:
        if 'session_id' not in session:
            return jsonify({'error': 'No session found'}), 400
        
        session_id = session['session_id']
        
        if session_id not in SESSIONS:
            return jsonify({'error': 'Session not found'}), 404
        
        # Find and remove the document
        original_count = len(SESSIONS[session_id]['documents'])
        SESSIONS[session_id]['documents'] = [
            doc for doc in SESSIONS[session_id]['documents'] 
            if doc['id'] != doc_id
        ]
        
        if len(SESSIONS[session_id]['documents']) < original_count:
            # Also try to delete the file
            for doc in SESSIONS[session_id]['documents']:
                if doc['id'] == doc_id and 'path' in doc:
                    try:
                        os.remove(doc['path'])
                    except:
                        pass  # File might already be deleted
            
            return jsonify({
                'success': True,
                'message': 'Document removed successfully',
                'document_count': len(SESSIONS[session_id]['documents'])
            })
        else:
            return jsonify({'error': 'Document not found'}), 404
            
    except Exception as e:
        logger.error(f"Remove document error: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/clear-messages', methods=['POST'])
def clear_messages():
    """Clear chat messages and reset token/cost counters from session."""
    try:
        if 'session_id' not in session:
            return jsonify({'error': 'No session found'}), 400
        
        session_id = session['session_id']
        if session_id in SESSIONS:
            # Clear messages
            SESSIONS[session_id]['messages'] = []
            # Reset token counts and costs
            SESSIONS[session_id]['total_tokens'] = {'input': 0, 'output': 0}
            SESSIONS[session_id]['total_cost'] = 0.0
            logger.info(f"Cleared messages and reset counters for session {session_id}")
            return jsonify({'success': True, 'message': 'Chat cleared'})
        else:
            return jsonify({'error': 'Session not found'}), 404
    except Exception as e:
        logger.error(f"Clear messages error: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/chat', methods=['POST'])
def chat():
    """Handle chat messages."""
    try:
        if 'session_id' not in session:
            return jsonify({'error': 'No session found'}), 400
        
        session_id = session['session_id']
        
        if session_id not in SESSIONS:
            SESSIONS[session_id] = {
                'documents': [],
                'messages': [],
                'total_tokens': {'input': 0, 'output': 0},
                'total_cost': 0.0,
                'selected_model': 'large',  # Default model
                'prompt_mode': 'stage1'  # Default prompt mode
            }
        
        data = request.json
        message = data.get('message', '')
        selected_model = data.get('model', SESSIONS[session_id].get('selected_model', 'large'))
        is_voice_mode = data.get('voice_mode', False)
        selected_document_ids = data.get('selected_documents', [])
        # Use 'voice' or 'text' prompt mode based on voice_mode flag
        prompt_mode = 'voice' if is_voice_mode else 'text'
        
        # Log model selection
        logger.info(f"Chat request - Model: {selected_model}, Model Name: {MODEL_OPTIONS[selected_model]['name']}, Voice Mode: {is_voice_mode}")
        
        if not message:
            return jsonify({'error': 'No message provided'}), 400
        
        # Update selected model and prompt mode
        SESSIONS[session_id]['selected_model'] = selected_model
        SESSIONS[session_id]['prompt_mode'] = prompt_mode
        
        # Add user message to history
        SESSIONS[session_id]['messages'].append({
            'role': 'user',
            'content': message
        })
        
        # Prepare conversation with documents
        conversation = {
            'messages': SESSIONS[session_id]['messages'],
            'model': selected_model,
            'prompt_mode': prompt_mode
        }
        
        # Add only SELECTED documents with metadata if available
        if selected_document_ids and SESSIONS[session_id]['documents']:
            # Filter to only include selected documents
            selected_docs = [
                doc for doc in SESSIONS[session_id]['documents']
                if doc['id'] in selected_document_ids
            ]
            if selected_docs:
                conversation['documents'] = selected_docs
                logger.info(f"Using {len(selected_docs)} selected documents out of {len(SESSIONS[session_id]['documents'])} total")
        
        # Stream response
        def generate():
            assistant_message = ""
            usage_info = None
            metrics_info = None
            chunk_count = 0
            last_content = ""
            
            try:
                for chunk in model(conversation):
                    chunk_count += 1
                    
                    if chunk.get('type') == 'assistant':
                        content = chunk.get('content', '')
                        last_content = content
                        assistant_message += content
                        
                        # Log every 10th chunk and any chunk with table markers
                        if chunk_count % 10 == 0 or '|' in content:
                            logger.debug(f"Chunk {chunk_count}: len={len(content)}, has_pipe={'|' in content}, preview={repr(content[:50])}")
                        
                        # Send as JSON for the frontend to parse
                        import json
                        try:
                            # Use ensure_ascii=False to handle special characters properly
                            json_data = json.dumps({'content': content}, ensure_ascii=False)
                            yield f"data: {json_data}\n\n"
                        except Exception as json_error:
                            logger.error(f"JSON encoding error at chunk {chunk_count}: {str(json_error)}, content={repr(content)}")
                            # Try to send with escaped content
                            escaped_content = content.encode('unicode_escape').decode('ascii')
                            yield f"data: {json.dumps({'content': escaped_content})}\n\n"
                            
                    elif chunk.get('type') == 'error':
                        error_msg = chunk.get('content', 'Unknown error')
                        logger.error(f"Error during streaming at chunk {chunk_count}: {error_msg}")
                        yield f"data: {json.dumps({'error': error_msg}, ensure_ascii=False)}\n\n"
                        
                    elif chunk.get('type') == 'usage':
                        # Capture usage information from the LLM response
                        usage_info = chunk.get('usage', {})
                        metrics_info = chunk.get('metrics', {})
                        
            except Exception as e:
                logger.error(f"Exception in generate function at chunk {chunk_count}: {str(e)}", exc_info=True)
                logger.error(f"Last content before error: {repr(last_content)}")
                logger.error(f"Total message so far ({len(assistant_message)} chars): {repr(assistant_message[:500])}")
                yield f"data: {json.dumps({'error': f'Streaming error at chunk {chunk_count}: {str(e)}'}, ensure_ascii=False)}\n\n"
            
            # Store assistant message and track tokens/cost from actual API response
            if assistant_message:
                SESSIONS[session_id]['messages'].append({
                    'role': 'assistant',
                    'content': assistant_message
                })
                
                # Update token counts and cost if we have usage info
                if usage_info:
                    input_tokens = usage_info.get('prompt_tokens', 0)
                    output_tokens = usage_info.get('completion_tokens', 0)
                    
                    SESSIONS[session_id]['total_tokens']['input'] += input_tokens
                    SESSIONS[session_id]['total_tokens']['output'] += output_tokens
                    
                    # Use the pre-calculated cost from llm_connector if available
                    if metrics_info and 'total_cost' in metrics_info:
                        SESSIONS[session_id]['total_cost'] += metrics_info['total_cost']
                    else:
                        # Fallback to manual calculation if metrics not available
                        model_config = MODEL_OPTIONS[selected_model]
                        input_cost = (input_tokens / 1000) * model_config['cost_input']
                        output_cost = (output_tokens / 1000) * model_config['cost_output']
                        SESSIONS[session_id]['total_cost'] += input_cost + output_cost
        
        return Response(generate(), mimetype='text/event-stream')
        
    except Exception as e:
        logger.error(f"Chat error: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/clear', methods=['POST'])
def clear_session():
    """Clear session documents and messages."""
    try:
        if 'session_id' in session:
            session_id = session['session_id']
            
            # Clean up uploaded files
            if session_id in SESSIONS:
                for doc in SESSIONS[session_id].get('documents', []):
                    if 'path' in doc and os.path.exists(doc['path']):
                        os.remove(doc['path'])
                
                # Clear session data
                SESSIONS[session_id] = {
                    'documents': [],
                    'messages': [],
                    'total_tokens': {'input': 0, 'output': 0},
                    'total_cost': 0.0,
                    'selected_model': 'large'  # Default model
                }
            
            logger.info(f"Session cleared: {session_id}")
        
        return jsonify({'success': True, 'message': 'Session cleared'})
        
    except Exception as e:
        logger.error(f"Clear session error: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/remove_document_old', methods=['POST'])
def remove_document_old():
    """Remove a specific document from the session (old endpoint)."""
    try:
        if 'session_id' not in session:
            return jsonify({'error': 'No session found'}), 400
        
        session_id = session['session_id']
        
        if session_id not in SESSIONS:
            return jsonify({'error': 'Session not found'}), 400
        
        data = request.json
        doc_id = data.get('document_id')
        
        if not doc_id:
            return jsonify({'error': 'No document_id provided'}), 400
        
        # Find and remove the document
        documents = SESSIONS[session_id]['documents']
        doc_to_remove = None
        for doc in documents:
            if doc['id'] == doc_id:
                doc_to_remove = doc
                break
        
        if doc_to_remove:
            # Remove file if it exists
            if 'path' in doc_to_remove and os.path.exists(doc_to_remove['path']):
                os.remove(doc_to_remove['path'])
            
            # Remove from session
            documents.remove(doc_to_remove)
            
            logger.info(f"Document removed: {doc_to_remove['filename']} from session {session_id}")
            
            return jsonify({
                'success': True,
                'message': f"Removed {doc_to_remove['filename']}",
                'documents': [{'id': d['id'], 'filename': d['filename']} for d in documents]
            })
        else:
            return jsonify({'error': 'Document not found'}), 404
            
    except Exception as e:
        logger.error(f"Remove document error: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/set_model', methods=['POST'])
def set_model():
    """Set the model for the session."""
    try:
        if 'session_id' not in session:
            return jsonify({'error': 'No session found'}), 400
        
        session_id = session['session_id']
        
        if session_id not in SESSIONS:
            SESSIONS[session_id] = {
                'documents': [],
                'messages': [],
                'total_tokens': {'input': 0, 'output': 0},
                'total_cost': 0.0,
                'selected_model': 'large',  # Default model
                'prompt_mode': 'stage1'  # Default prompt mode
            }
        
        data = request.json
        model_size = data.get('model')
        
        if model_size not in MODEL_OPTIONS:
            return jsonify({'error': 'Invalid model size'}), 400
        
        SESSIONS[session_id]['selected_model'] = model_size
        
        return jsonify({
            'success': True,
            'model': model_size,
            'model_name': MODEL_OPTIONS[model_size]['name']
        })
        
    except Exception as e:
        logger.error(f"Set model error: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/set_prompt', methods=['POST'])
def set_prompt():
    """Set the prompt mode for the session."""
    try:
        if 'session_id' not in session:
            return jsonify({'error': 'No session found'}), 400
        
        session_id = session['session_id']
        
        if session_id not in SESSIONS:
            SESSIONS[session_id] = {
                'documents': [],
                'messages': [],
                'total_tokens': {'input': 0, 'output': 0},
                'total_cost': 0.0,
                'selected_model': 'large',  # Default model
                'prompt_mode': 'stage1'  # Default prompt mode
            }
        
        data = request.json
        prompt_mode = data.get('prompt_mode')
        
        valid_modes = ['stage1', 'stage2', 'basic', 'default']
        if prompt_mode not in valid_modes:
            return jsonify({'error': 'Invalid prompt mode'}), 400
        
        SESSIONS[session_id]['prompt_mode'] = prompt_mode
        
        return jsonify({
            'success': True,
            'prompt_mode': prompt_mode
        })
        
    except Exception as e:
        logger.error(f"Set prompt error: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/models', methods=['GET'])
def get_models():
    """Get available models."""
    return jsonify(MODEL_OPTIONS)


@app.route('/status', methods=['GET'])
def status():
    """Get current session status."""
    try:
        if 'session_id' not in session:
            return jsonify({
                'session_id': None,
                'document_count': 0,
                'message_count': 0
            })
        
        session_id = session['session_id']
        
        if session_id not in SESSIONS:
            SESSIONS[session_id] = {
                'documents': [],
                'messages': [],
                'total_tokens': {'input': 0, 'output': 0},
                'total_cost': 0.0,
                'selected_model': 'large',  # Default model
                'prompt_mode': 'stage1'  # Default prompt mode
            }
        
        return jsonify({
            'session_id': session_id,
            'document_count': len(SESSIONS[session_id]['documents']),
            'message_count': len(SESSIONS[session_id]['messages']),
            'documents': [
                {
                    'id': doc['id'],
                    'filename': doc['filename'],
                    'size': os.path.getsize(doc['path']) if os.path.exists(doc['path']) else 0,
                    'token_count': len(doc.get('content', '').split())
                }
                for doc in SESSIONS[session_id]['documents']
            ],
            'selected_model': SESSIONS[session_id].get('selected_model', 'large'),
            'total_tokens': SESSIONS[session_id].get('total_tokens', {'input': 0, 'output': 0}),
            'total_cost': round(SESSIONS[session_id].get('total_cost', 0.0), 4)
        })
        
    except Exception as e:
        logger.error(f"Status error: {e}")
        return jsonify({'error': str(e)}), 500


# ====================
# Voice API Endpoints 
# ====================

@app.route('/transcribe', methods=['POST'])
def transcribe_audio():
    """Transcribe audio using Whisper STT."""
    try:
        # Check if audio file is in the request
        if 'audio' not in request.files:
            return jsonify({"error": "No audio file provided"}), 400
        
        audio_file = request.files['audio']
        
        # Get the model parameter from form data
        model_size = request.form.get('model', default_whisper_model)
        if model_size not in WHISPER_MODELS:
            model_size = default_whisper_model
        
        # Medium model is now downloaded and available
        # Large model may still need downloading
        if model_size == 'large':
            print(f"Model {model_size} may need downloading, falling back to {default_whisper_model}")
            model_size = default_whisper_model
        
        selected_model_path = WHISPER_MODELS[model_size]
        print(f"Using Whisper model: {model_size} ({selected_model_path})")
        
        # Check if file is empty
        audio_file.seek(0, os.SEEK_END)
        file_size = audio_file.tell()
        audio_file.seek(0)
        
        if file_size == 0:
            print("Received empty audio file")
            return jsonify({"error": "Audio file is empty"}), 400
        
        print(f"Received audio file, size: {file_size} bytes")
        
        # Save the uploaded file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix='.webm') as tmp_file:
            audio_file.save(tmp_file.name)
            temp_path = tmp_file.name
        
        try:
            print(f"Transcribing audio file: {temp_path}")
            start_time = time.time()
            
            # Check file size on disk
            disk_size = os.path.getsize(temp_path)
            print(f"Saved file size: {disk_size} bytes")
            
            if disk_size < 100:  # Too small to be valid audio
                print("Audio file too small to process")
                return jsonify({"error": "Audio file too small"}), 400
            
            # Transcribe the audio with selected model
            result = mlx_whisper.transcribe(
                temp_path,
                path_or_hf_repo=selected_model_path,
                verbose=False
            )
            
            transcription = result["text"].strip()
            
            transcription_time = time.time() - start_time
            print(f"Transcription complete in {transcription_time:.2f}s: {transcription}")
            
            if not transcription:
                return jsonify({"error": "No speech detected"}), 400
            
            return jsonify({
                "text": transcription,
                "time": transcription_time
            })
            
        finally:
            # Clean up temporary file
            if os.path.exists(temp_path):
                os.remove(temp_path)
        
    except Exception as e:
        print(f"Error transcribing audio: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


@app.route('/generate', methods=['POST'])
def generate_audio():
    """Generate audio using Kokoro TTS."""
    try:
        data = request.json
        text = data.get('text', '')
        # Accept voice and speed parameters from client
        voice = data.get('voice', 'af_aoede')
        speed = data.get('speed', 1.27)
        
        if not text:
            return jsonify({"error": "No text provided"}), 400
        
        print(f"Generating audio for: {text} (voice: {voice}, speed: {speed})")
        start_time = time.time()
        
        # Generate audio using Kokoro with specified voice
        # The pipeline returns a generator of (graphemes, phonemes, audio)
        result = list(tts_pipeline(text, voice=voice, speed=speed))
        if result:
            _, _, audio_array = result[0]
        else:
            raise ValueError("No audio generated")
        
        # Convert to numpy array if needed
        if hasattr(audio_array, 'numpy'):
            audio_array = audio_array.numpy()
        
        # Ensure audio is in the right format (float32, -1 to 1 range)
        audio_array = np.array(audio_array, dtype=np.float32)
        
        # If audio is 2D, take the first channel
        if len(audio_array.shape) > 1:
            audio_array = audio_array[0]
        
        # Normalize if needed
        max_val = np.max(np.abs(audio_array))
        if max_val > 1.0:
            audio_array = audio_array / max_val
        
        generation_time = time.time() - start_time
        print(f"Audio generated in {generation_time:.2f} seconds")
        
        # Convert to WAV bytes
        buffer = io.BytesIO()
        sf.write(buffer, audio_array, samplerate=24000, format='WAV')
        buffer.seek(0)
        
        return Response(
            buffer.getvalue(),
            mimetype='audio/wav',
            headers={
                'Content-Disposition': 'inline; filename="output.wav"',
                'X-Generation-Time': str(generation_time)
            }
        )
        
    except Exception as e:
        print(f"Error generating audio: {str(e)}")
        return jsonify({"error": str(e)}), 500


if __name__ == '__main__':
    print("\n" + "="*50)
    print("✅ Unified Chat+ Voice Server Starting")
    print("="*50)
    print("\n🌐 Access the application at:")
    print("   Regular chat: http://localhost:5003")
    print("   Voice chat:   http://localhost:5003?voice=true")
    print("\n✨ All voice capabilities integrated into single server!")
    print("="*50 + "\n")
    app.run(debug=True, port=5003)